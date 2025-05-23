from flask import Flask, render_template, request, redirect, url_for, session, send_file
from werkzeug.security import generate_password_hash, check_password_hash
import sqlite3
from datetime import datetime
import calendar
import pandas as pd
import tempfile
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl import load_workbook
import locale
from datetime import datetime
from openpyxl.utils import get_column_letter
import tempfile
from flask import send_file
import locale
from datetime import datetime
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from collections import defaultdict
from flask import request, send_file
from datetime import datetime
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from collections import defaultdict
import tempfile
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT




app = Flask(__name__)
app.secret_key = 'chave_secreta'

def init_db():
    with sqlite3.connect("usuarios.db") as conn:
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS usuarios (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        cargo TEXT,
                        nome TEXT,
                        email TEXT UNIQUE,
                        senha TEXT,
                        telefone TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS opcoes_lista (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        opcao TEXT)''')
        conn.commit()

@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form["email"]
        senha = request.form["senha"]
        with sqlite3.connect("usuarios.db") as conn:
            c = conn.cursor()
            c.execute("SELECT * FROM usuarios WHERE email = ?", (email,))
            user = c.fetchone()
            if user and check_password_hash(user[4], senha):
                session["user_id"] = user[0]
                session["nome"] = user[2]
                return redirect(url_for("painel"))
            else:
                return "Usuário ou senha inválidos"
    return render_template("login.html")

@app.route("/cadastro", methods=["GET", "POST"])
def cadastro():
    if request.method == "POST":
        cargo = request.form["cargo"]
        nome = request.form["nome"]
        email = request.form["email"]
        senha = request.form["senha"]
        confirmacao = request.form["confirmacao"]
        telefone = request.form.get("telefone", "")

        if senha != confirmacao:
            return "As senhas não coincidem!"

        senha_hash = generate_password_hash(senha)

        try:
            with sqlite3.connect("usuarios.db") as conn:
                c = conn.cursor()
                c.execute("INSERT INTO usuarios (cargo, nome, email, senha, telefone) VALUES (?, ?, ?, ?, ?)",
                          (cargo, nome, email, senha_hash, telefone))
                conn.commit()
            return redirect(url_for("login"))
        except sqlite3.IntegrityError:
            return "E-mail já cadastrado!"

    return render_template("cadastro.html")

import locale
from datetime import datetime
import calendar

@app.route("/painel", methods=["GET"])
def painel():
    if "user_id" not in session:
        return redirect(url_for("login"))

    import locale
    import platform
    from datetime import datetime
    import calendar
    import sqlite3

    try:
        if platform.system() == 'Windows':
            locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')
        else:
            locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
    except locale.Error:
        print("Aviso: Locale português não pôde ser configurado. Datas podem aparecer em inglês.")

    hoje = datetime.now()

    # Calcular o próximo mês e ano correspondente
    if hoje.month == 12:
        proximo_mes = 1
        ano = hoje.year + 1
    else:
        proximo_mes = hoje.month + 1
        ano = hoje.year

    dias_mes = calendar.monthrange(ano, proximo_mes)[1]

    dias_semana_pt = {
        "Monday": "Segunda-feira",
        "Tuesday": "Terça-feira",
        "Wednesday": "Quarta-feira",
        "Thursday": "Quinta-feira",
        "Friday": "Sexta-feira",
        "Saturday": "Sábado",
        "Sunday": "Domingo"
    }

    datas = []
    for d in range(1, dias_mes + 1):
        data_obj = datetime(ano, proximo_mes, d)
        dia_str = f"{ano}-{proximo_mes:02d}-{d:02d}"
        dia_semana_eng = data_obj.strftime('%A')
        dia_semana = dias_semana_pt.get(dia_semana_eng, dia_semana_eng)
        datas.append({
            "dia": dia_str,
            "dia_semana": dia_semana
        })

    with sqlite3.connect("usuarios.db") as conn:
        c = conn.cursor()
        c.execute("SELECT opcao FROM opcoes_lista")
        opcoes = [row[0] for row in c.fetchall()]

    meses_pt = {
        "January": "Janeiro",
        "February": "Fevereiro",
        "March": "Março",
        "April": "Abril",
        "May": "Maio",
        "June": "Junho",
        "July": "Julho",
        "August": "Agosto",
        "September": "Setembro",
        "October": "Outubro",
        "November": "Novembro",
        "December": "Dezembro"
    }

    nome_mes_ingles = datetime(ano, proximo_mes, 1).strftime('%B')
    nome_mes = meses_pt.get(nome_mes_ingles, nome_mes_ingles)

    return render_template("painel.html", datas=datas, opcoes=opcoes, nome_mes=nome_mes)

@app.route("/gerar_excel", methods=["POST"])
def gerar_excel():
    dados = request.get_json()
    if not dados:
        return "Dados inválidos", 400

    hoje = datetime.now()

    meses_pt = {
        "January": "Janeiro", "February": "Fevereiro", "March": "Março", "April": "Abril",
        "May": "Maio", "June": "Junho", "July": "Julho", "August": "Agosto",
        "September": "Setembro", "October": "Outubro", "November": "Novembro", "December": "Dezembro"
    }
    nome_mes = meses_pt.get(hoje.strftime('%B'), hoje.strftime('%B'))

    dias_semana_pt = {
        "Monday": "Segunda-feira", "Tuesday": "Terça-feira", "Wednesday": "Quarta-feira",
        "Thursday": "Quinta-feira", "Friday": "Sexta-feira", "Saturday": "Sábado", "Sunday": "Domingo"
    }

    agrupados = defaultdict(list)
    for item in dados:
        try:
            dt = datetime.strptime(item.get('dia', ''), '%Y-%m-%d')
            data_formatada = dt.strftime('%d-%m-%Y')
            dia_semana_pt_nome = dias_semana_pt.get(dt.strftime('%A'), dt.strftime('%A'))
        except Exception:
            data_formatada = item.get('dia', '')
            dia_semana_pt_nome = item.get('dia_semana', '')

        data_completa = f"{data_formatada}\n{dia_semana_pt_nome}"

        atividades_raw = item.get('atividade', [])
        if isinstance(atividades_raw, str):
            atividades_raw = [atividades_raw]

        atividades_formatadas = [f"• {a}" for a in atividades_raw if a]

        horario = ""
        if item.get("hora_inicio") or item.get("hora_fim"):
            horario = f"{item.get('hora_inicio', '')} às {item.get('hora_fim', '')}".strip(" -")

        for atv in atividades_formatadas:
            agrupados[data_completa].append(f"    {atv}   (HORÁRIO:  {horario})" if horario else f"    {atv}")

    dados_final = []
    for data_completa, lista_atividades in agrupados.items():
        atividades_concatenadas = "\n".join(lista_atividades)
        dados_final.append({
            "Data": data_completa,
            "Atividades - Horário": atividades_concatenadas
        })

    df = pd.DataFrame(dados_final)[['Data', 'Atividades - Horário']]

    # === Excel ===
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp_excel:
        caminho_excel = tmp_excel.name

    df.to_excel(caminho_excel, index=False)
    wb = load_workbook(caminho_excel)
    ws = wb.active
    ws.insert_rows(1, amount=12)
    ws.insert_cols(1, amount=3)

    start_row = 12
    start_col = 4

    cabecalho = [
        "PARÓQUIA NOSSA SENHORA DAS GRAÇAS",
        "ARQUIDIOCESE DE MARINGÁ – PR",
        "Pároco: Pe. Leomar Antônio Montagna – Vigário: Pe. Paulo Felipe dos Santos",
        "Diáconos: Antônio Bueno de Camargo e Carlos Roberto Paulino",
        "Praça Ipiranga, 271, centro - Sarandi – PR",
        "CEP: 87111-005 - Fone: (44) 3035-3011 / 99750 5254",
        "Facebook: paroquia nossa senhora das graças",
        "E-mail: paroquiasdi@gmail.com"
    ]

    for i, linha in enumerate(cabecalho, start=2):
        ws.merge_cells(start_row=i, start_column=start_col, end_row=i, end_column=start_col + 1)
        cell = ws.cell(row=i, column=start_col)
        cell.value = linha
        cell.font = Font(size=11, bold=(i == 2))
        cell.alignment = Alignment(horizontal='center')

    titulo = f"AGENDA MENSAL IGREJA NOSSA SRA. DAS GRAÇAS - Mês de {nome_mes}"
    ws.merge_cells(start_row=10, start_column=start_col, end_row=10, end_column=start_col + 1)
    cell_titulo = ws.cell(row=10, column=start_col)
    cell_titulo.value = titulo
    cell_titulo.font = Font(size=14, bold=True)
    cell_titulo.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[10].height = 35

    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    for i, col in enumerate(range(start_col, start_col + 2)):
        cell = ws.cell(row=start_row, column=col)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        ws.column_dimensions[get_column_letter(col)].width = [25, 100][i]
    ws.row_dimensions[start_row].height = 40

    thin_border = Border(
        left=Side(style="thin", color="A9B7C6"),
        right=Side(style="thin", color="A9B7C6"),
        top=Side(style="thin", color="A9B7C6"),
        bottom=Side(style="thin", color="A9B7C6")
    )

    for row in range(start_row + 1, ws.max_row + 1):
        fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid") if (row - start_row) % 2 else None
        for col in range(start_col, start_col + 2):
            cell = ws.cell(row=row, column=col)
            cell.border = thin_border
            if col == start_col:
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            else:
                cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
            if fill:
                cell.fill = fill
            ws.row_dimensions[row].height = 80

    wb.save(caminho_excel)

    # === Word ===
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_doc:
        caminho_doc = tmp_doc.name

    doc = Document()

    for linha in cabecalho:
        p = doc.add_paragraph(linha)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.style.font.size = Pt(10)
        if "PARÓQUIA" in linha:
            p.style.font.bold = True

    doc.add_paragraph()  # Espaço

    doc.add_paragraph(titulo).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for item in dados_final:
        doc.add_paragraph()
        p_data = doc.add_paragraph(item["Data"])
        p_data.runs[0].bold = True
        doc.add_paragraph(item["Atividades - Horário"])

    doc.save(caminho_doc)

    # Envie os dois arquivos (Excel e Word) — como exemplo: somente Excel aqui
    return send_file(caminho_excel, as_attachment=True, download_name="agenda_mensal.xlsx")
    # Para enviar o Word: return send_file(caminho_doc, as_attachment=True, download_name="agenda_mensal.docx")
@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", port=5000, debug=True)
